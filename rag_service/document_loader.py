import json
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable, List, Optional, Tuple, Union

from .models import Document

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".html", ".ipynb"}


PathLike = Union[str, Path]
DocumentInput = Union[PathLike, Tuple[PathLike, str]]


def load_document(path: PathLike, source_name: Optional[str] = None) -> Document:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {supported}")

    if suffix == ".pdf":
        text = _read_pdf(file_path)
    elif suffix == ".docx":
        text = _read_docx(file_path)
    elif suffix == ".html":
        text = _read_html(file_path)
    elif suffix == ".ipynb":
        text = _read_ipynb(file_path)
    else:
        text = _read_text_file(file_path)

    return Document(
        text=_normalize_text(text),
        metadata={
            "source": source_name or file_path.name,
            "path": str(file_path),
            "extension": suffix,
        },
    )


def load_documents(paths: Iterable[DocumentInput]) -> List[Document]:
    documents = []
    for item in paths:
        if isinstance(item, tuple):
            documents.extend(load_document_blocks(item[0], source_name=item[1]))
        else:
            documents.extend(load_document_blocks(item))
    return documents


def load_document_blocks(path: PathLike, source_name: Optional[str] = None) -> List[Document]:
    """Load one or more retrievable blocks from a file.

    PDF files are split into one document per page. Other supported files are
    returned as a single document.
    """
    file_path = Path(path)
    if file_path.suffix.lower() == ".pdf":
        return _read_pdf_pages(file_path, source_name=source_name)
    return [load_document(file_path, source_name=source_name)]


def _read_pdf_pages(path: Path, source_name: Optional[str] = None) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    documents = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            documents.append(
                Document(
                    text=_normalize_text(page_text),
                    metadata={
                        "source": source_name or path.name,
                        "path": str(path),
                        "extension": ".pdf",
                        "page_number": page_number,
                    },
                )
            )
    return documents


def _read_pdf(path: Path) -> str:
    pages = _read_pdf_pages(path)
    return "\n\n".join(f"[page {page.metadata['page_number']}]\n{page.text}" for page in pages)


def _read_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            table_cells.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs + table_cells)


def _read_html(path: Path) -> str:
    parser = _TextHTMLParser()
    parser.feed(_read_text_file(path))
    return parser.text


def _read_ipynb(path: Path) -> str:
    notebook = json.loads(_read_text_file(path))
    blocks = []
    for cell in notebook.get("cells", []):
        cell_type = cell.get("cell_type", "cell")
        source = cell.get("source", "")
        if isinstance(source, list):
            source = "".join(source)
        if source.strip():
            blocks.append(f"[{cell_type}]\n{source}")
    return "\n\n".join(blocks)


def _read_text_file(path: Path) -> str:
    last_error = None
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as error:
            last_error = error
    raise last_error


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    normalized_lines = [line for line in lines if line]
    return "\n".join(normalized_lines)


class _TextHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []

    @property
    def text(self) -> str:
        return "\n".join(self.parts)

    def handle_data(self, data: str) -> None:
        text = data.strip()
        if text:
            self.parts.append(text)
